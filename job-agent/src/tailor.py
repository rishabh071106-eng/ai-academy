"""Resume + cover letter tailoring.

Claude rewrites the master resume for a specific job (re-weighting emphasis,
never inventing facts), returns structured content, and we render it to DOCX.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from .config import Config
from .discover import Job
from .llm import LLM

RESUME_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["headline", "summary", "competencies", "experience",
                 "education", "certifications", "cover_letter"],
    "properties": {
        "headline": {"type": "string", "description": "One-line positioning under the name"},
        "summary": {"type": "string", "description": "4-6 sentence professional summary tailored to the job"},
        "competencies": {
            "type": "array", "items": {"type": "string"},
            "description": "10-14 short competency phrases, most job-relevant first",
        },
        "experience": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["title", "company", "period", "bullets"],
                "properties": {
                    "title": {"type": "string"},
                    "company": {"type": "string"},
                    "period": {"type": "string"},
                    "bullets": {"type": "array", "items": {"type": "string"}},
                },
            },
        },
        "education": {"type": "array", "items": {"type": "string"}},
        "certifications": {"type": "array", "items": {"type": "string"}},
        "cover_letter": {
            "type": "string",
            "description": "250-350 word cover letter addressed to the hiring team",
        },
    },
}


def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")[:60]


def tailor(job: Job, emphasis: str, cfg: Config, llm: LLM) -> dict:
    prompt = f"""Tailor this candidate's resume and write a cover letter for the job below.

MASTER RESUME (single source of truth — you may rephrase, reorder, re-weight,
and omit, but NEVER invent employers, titles, dates, metrics, or credentials
that are not here):
{cfg.resume_master}

JOB:
Company: {job.company}
Title: {job.title}
Location: {job.location}
Description:
{job.description[:8000]}

EMPHASIS: {emphasis} (see the master resume's tailoring guidance).

Requirements:
- Mirror the job description's language where truthful (helps ATS keyword matching).
- Lead every experience bullet with impact; keep the strongest, most relevant
  6-8 bullets for JP Morgan and 4-5 for Wipro.
- The summary must read as a {job.title} candidate, not a generic PM.
- Cover letter: specific to this company and role, confident, no flattery filler,
  mention the candidate is a current VP of Product at a global bank leading
  AI-native platform transformation."""

    return llm.ask_json(prompt, RESUME_SCHEMA, effort="high", max_tokens=12000)


def render_docx(content: dict, profile: dict, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    name = doc.add_paragraph()
    run = name.add_run(profile["name"])
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph(content["headline"])
    doc.add_paragraph(
        f'{profile["location"]} · {profile["phone"]} · {profile["email"]} · {profile["linkedin"]}'
    )

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)

    heading("Summary")
    doc.add_paragraph(content["summary"])

    heading("Core Competencies")
    doc.add_paragraph(" · ".join(content["competencies"]))

    heading("Professional Experience")
    for role in content["experience"]:
        p = doc.add_paragraph()
        r = p.add_run(f'{role["title"]} | {role["company"]}')
        r.bold = True
        doc.add_paragraph(role["period"]).runs[0].italic = True
        for bullet in role["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")

    heading("Education")
    for item in content["education"]:
        doc.add_paragraph(item, style="List Bullet")

    heading("Certifications")
    for item in content["certifications"]:
        doc.add_paragraph(item, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_application_package(job: Job, emphasis: str, cfg: Config, llm: LLM) -> Path:
    """Returns the directory containing resume.docx + cover_letter.txt."""
    content = tailor(job, emphasis, cfg, llm)
    pkg_dir = cfg.path("applications_dir", "applications") / f"{slugify(job.company)}--{slugify(job.title)}"
    render_docx(content, cfg.profile, pkg_dir / "resume.docx")
    (pkg_dir / "cover_letter.txt").write_text(content["cover_letter"])
    (pkg_dir / "job.txt").write_text(
        f"{job.title} @ {job.company}\n{job.url}\n{job.location}\n\n{job.description[:4000]}"
    )
    return pkg_dir
